from io import BytesIO

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.views.decorators.http import require_POST
from django.db.models import Q, Count
from django.http import FileResponse
from django.core.files.base import ContentFile

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Document, Client, ClientCase, RequiredDocument
from .forms import DocumentForm, StaffClientCaseForm, ClientCaseEditForm, FeedbackForm


def is_staff_user(user):
    return user.is_authenticated and user.is_staff


@login_required
def after_login_redirect(request):
    if request.user.is_staff:
        return redirect('staff_dashboard')

    return redirect('document_list')


@login_required
def document_list(request):
    if request.user.is_staff:
        required_documents = RequiredDocument.objects.all().select_related(
            'client_case',
            'client_case__client',
            'client_case__service_type',
            'document_type'
        )
    else:
        required_documents = RequiredDocument.objects.filter(
            client_case__client__user=request.user
        ).select_related(
            'client_case',
            'client_case__client',
            'client_case__service_type',
            'document_type'
        )

    search_query = request.GET.get('q')

    if search_query:
        required_documents = required_documents.filter(
            Q(client_case__client__full_name__icontains=search_query) |
            Q(client_case__client__inn__icontains=search_query) |
            Q(document_type__name__icontains=search_query) |
            Q(client_case__service_type__name__icontains=search_query)
        )

    context = {
        'required_documents': required_documents,
        'search_query': search_query,
        'status_choices': RequiredDocument.STATUS_CHOICES,
    }

    return render(request, 'documents/document_list.html', context)


@login_required
def document_detail(request, pk):
    document = get_object_or_404(Document, pk=pk)

    if not request.user.is_staff and document.required_document.client_case.client.user != request.user:
        return redirect('document_list')

    next_url = request.GET.get('next')

    return render(request, 'documents/document_detail.html', {
        'document': document,
        'next_url': next_url,
    })


@login_required
def document_create(request):
    required_document_id = request.GET.get('required_document')
    required_document = get_object_or_404(RequiredDocument, pk=required_document_id)

    if not request.user.is_staff and required_document.client_case.client.user != request.user:
        return redirect('document_list')

    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.required_document = required_document
            document.save()

            required_document.status = 'uploaded'
            required_document.save()

            return redirect('document_list')
    else:
        form = DocumentForm()

    return render(request, 'documents/document_form.html', {
        'form': form,
        'required_document': required_document,
    })


@login_required
def document_file(request, pk):
    document = get_object_or_404(Document, pk=pk)

    if not request.user.is_staff and document.required_document.client_case.client.user != request.user:
        return redirect('document_list')

    download = request.GET.get('download') == '1'

    return FileResponse(
        document.file.open('rb'),
        as_attachment=download,
        filename=document.file.name.split('/')[-1]
    )


@login_required
@require_POST
def required_document_status_update(request, pk):
    if not request.user.is_staff:
        return redirect('document_list')

    required_document = get_object_or_404(RequiredDocument, pk=pk)
    new_status = request.POST.get('status')

    allowed_statuses = [
        'brought_personally',
        'not_required',
        'verified',
        'rejected',
        'waiting',
    ]

    if new_status in allowed_statuses:
        required_document.status = new_status
        required_document.save()

    next_url = request.POST.get('next')

    if next_url:
        return redirect(next_url)

    return redirect('staff_client_case_detail', pk=required_document.client_case.pk)


@login_required
@user_passes_test(is_staff_user)
def staff_dashboard(request):
    search_query = request.GET.get('q')

    cases = ClientCase.objects.all().select_related(
        'client',
        'service_type'
    ).annotate(
        total_documents=Count('required_documents'),
        completed_documents=Count(
            'required_documents',
            filter=Q(required_documents__status__in=[
                'uploaded',
                'verified',
                'brought_personally',
                'not_required',
            ])
        ),
        waiting_documents=Count(
            'required_documents',
            filter=Q(required_documents__status='waiting')
        )
    ).order_by('-created_at')

    if search_query:
        cases = cases.filter(
            Q(client__full_name__icontains=search_query) |
            Q(client__inn__icontains=search_query) |
            Q(service_type__name__icontains=search_query)
        )

    context = {
        'cases': cases,
        'search_query': search_query,
    }

    return render(request, 'documents/staff_dashboard.html', context)


@login_required
@user_passes_test(is_staff_user)
def staff_client_case_create(request):
    if request.method == 'POST':
        form = StaffClientCaseForm(request.POST)

        if form.is_valid():
            inn = form.cleaned_data['inn']
            full_name = form.cleaned_data['full_name']
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            service_type = form.cleaned_data['service_type']

            organization_name = form.cleaned_data.get('organization_name')
            short_organization_name = form.cleaned_data.get('short_organization_name')
            legal_address = form.cleaned_data.get('legal_address')
            director_full_name = form.cleaned_data.get('director_full_name')
            authorized_capital = form.cleaned_data.get('authorized_capital')

            client = Client.objects.filter(inn=inn).first()

            if client is None:
                user = User.objects.create_user(
                    username=username,
                    password=password
                )

                client = Client.objects.create(
                    user=user,
                    full_name=full_name,
                    inn=inn
                )

            ClientCase.objects.create(
                client=client,
                service_type=service_type,
                organization_name=organization_name,
                short_organization_name=short_organization_name,
                legal_address=legal_address,
                director_full_name=director_full_name,
                authorized_capital=authorized_capital,
            )

            return redirect('staff_dashboard')
    else:
        form = StaffClientCaseForm()

    return render(request, 'documents/staff_client_case_create.html', {
        'form': form
    })


@login_required
@user_passes_test(is_staff_user)
def generate_charter(request, pk):
    required_document = get_object_or_404(RequiredDocument, pk=pk)

    if required_document.document_type.name.lower() != 'устав':
        return redirect('document_list')

    client_case = required_document.client_case
    client = client_case.client
    service_type = client_case.service_type

    def value_or_placeholder(value):
        return value if value else '[не указано]'

    organization_name = value_or_placeholder(client_case.organization_name)
    short_organization_name = value_or_placeholder(client_case.short_organization_name)
    legal_address = value_or_placeholder(client_case.legal_address)
    director_full_name = value_or_placeholder(client_case.director_full_name)
    authorized_capital = value_or_placeholder(client_case.authorized_capital)

    doc = DocxDocument()

    # Настройка базового шрифта
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Титульная часть
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('УТВЕРЖДЁН\n')
    run.bold = True
    p.add_run('решением учредителя\n')
    p.add_run('от «___» __________ 2026 г.')

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('УСТАВ')
    run.bold = True
    run.font.size = Pt(16)

    org_title = doc.add_paragraph()
    org_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_title.add_run(organization_name)
    run.bold = True

    if short_organization_name != '[не указано]':
        short_title = doc.add_paragraph()
        short_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        short_title.add_run(f'({short_organization_name})')

    doc.add_paragraph()

    city = doc.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city.add_run('г. Казань\n2026 г.')

    doc.add_page_break()

    # Раздел 1
    doc.add_heading('1. Общие положения', level=1)
    doc.add_paragraph(
        f'1.1. {organization_name}, именуемая в дальнейшем «Организация», '
        f'создаётся в рамках услуги: {service_type.name}.'
    )
    doc.add_paragraph(
        f'1.2. Сокращённое наименование организации: {short_organization_name}.'
    )
    doc.add_paragraph(
        f'1.3. Место нахождения и юридический адрес организации: {legal_address}.'
    )
    doc.add_paragraph(
        '1.4. Организация осуществляет свою деятельность в соответствии с '
        'законодательством Российской Федерации, настоящим уставом и внутренними документами.'
    )

    # Раздел 2
    doc.add_heading('2. Учредитель и сведения о клиенте', level=1)
    doc.add_paragraph(f'2.1. Учредитель: {client.full_name}.')
    doc.add_paragraph(f'2.2. ИНН учредителя: {client.inn}.')
    doc.add_paragraph(f'2.3. Руководитель организации: {director_full_name}.')

    # Раздел 3
    doc.add_heading('3. Цели и виды деятельности', level=1)
    doc.add_paragraph(
        '3.1. Организация создаётся для осуществления деятельности, не запрещённой '
        'законодательством Российской Федерации.'
    )
    doc.add_paragraph(
        '3.2. Конкретные виды деятельности определяются учредителем и могут быть уточнены '
        'при подготовке итогового пакета документов.'
    )

    # Раздел 4
    doc.add_heading('4. Имущество и уставный капитал', level=1)
    doc.add_paragraph(
        f'4.1. Размер уставного капитала / имущественной базы: {authorized_capital}.'
    )
    doc.add_paragraph(
        '4.2. Источники формирования имущества определяются в соответствии с '
        'законодательством Российской Федерации и решениями учредителя.'
    )

    # Раздел 5
    doc.add_heading('5. Управление организацией', level=1)
    doc.add_paragraph(
        f'5.1. Руководство текущей деятельностью организации осуществляет: {director_full_name}.'
    )
    doc.add_paragraph(
        '5.2. Органы управления, их компетенция и порядок принятия решений определяются '
        'законодательством Российской Федерации и настоящим уставом.'
    )

    # Раздел 6
    doc.add_heading('6. Документы организации', level=1)
    doc.add_paragraph(
        '6.1. Организация обеспечивает хранение учредительных, регистрационных, бухгалтерских '
        'и иных документов в порядке, установленном законодательством Российской Федерации.'
    )
    doc.add_paragraph(
        '6.2. Электронные копии документов могут храниться в информационной системе '
        'автоматизированной обработки и хранения электронных документов.'
    )

    # Раздел 7
    doc.add_heading('7. Реорганизация и ликвидация', level=1)
    doc.add_paragraph(
        '7.1. Реорганизация и ликвидация организации осуществляются в порядке, установленном '
        'законодательством Российской Федерации.'
    )

    # Заключение
    doc.add_heading('8. Заключительные положения', level=1)
    doc.add_paragraph(
        '8.1. Настоящий документ является автоматически сформированным черновым вариантом '
        'и подлежит проверке сотрудником перед использованием в официальном пакете документов.'
    )
    doc.add_paragraph(
        '8.2. Поля со значением «[не указано]» должны быть дополнены сотрудником до передачи '
        'документа на дальнейшее оформление.'
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"ustav_{client.inn}.docx"

    # Если у этого пункта уже был файл — создаём новый файл,
    # старый останется в истории, но в карточке откроется первый.
    generated_document = Document.objects.create(
        required_document=required_document
    )

    generated_document.file.save(
        filename,
        ContentFile(buffer.read()),
        save=True
    )

    required_document.status = 'uploaded'
    required_document.save()

    return redirect('document_list')

@login_required
@user_passes_test(is_staff_user)
def staff_client_case_edit(request, pk):
    client_case = get_object_or_404(ClientCase, pk=pk)

    if request.method == 'POST':
        form = ClientCaseEditForm(request.POST, instance=client_case)

        if form.is_valid():
            form.save()
            return redirect('staff_dashboard')
    else:
        form = ClientCaseEditForm(instance=client_case)

    return render(request, 'documents/staff_client_case_edit.html', {
        'form': form,
        'client_case': client_case,
    })

@login_required
@require_POST
def document_delete(request, pk):
    document = get_object_or_404(Document, pk=pk)
    required_document = document.required_document

    if not request.user.is_staff and required_document.client_case.client.user != request.user:
        return redirect('document_list')

    if document.file:
        document.file.delete(save=False)

    document.delete()

    if not required_document.files.exists():
        required_document.status = 'waiting'
        required_document.save()

    return redirect('document_list')

@login_required
def document_replace(request, pk):
    document = get_object_or_404(Document, pk=pk)
    required_document = document.required_document

    if not request.user.is_staff and required_document.client_case.client.user != request.user:
        return redirect('document_list')

    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)

        if form.is_valid():
            new_file = form.cleaned_data['file']

            if document.file:
                document.file.delete(save=False)

            document.file.save(
                new_file.name,
                new_file,
                save=True
            )

            required_document.status = 'uploaded'
            required_document.save()

            return redirect('document_detail', pk=document.pk)
    else:
        form = DocumentForm()

    return render(request, 'documents/document_replace.html', {
        'form': form,
        'document': document,
    })

@login_required
@user_passes_test(is_staff_user)
def staff_client_case_detail(request, pk):
    client_case = get_object_or_404(
        ClientCase.objects.select_related('client', 'service_type'),
        pk=pk
    )

    required_documents = RequiredDocument.objects.filter(
        client_case=client_case
    ).select_related(
        'document_type',
        'client_case',
        'client_case__client',
        'client_case__service_type'
    )

    total_documents = required_documents.count()
    completed_documents = required_documents.filter(
        status__in=[
            'uploaded',
            'verified',
            'brought_personally',
            'not_required',
        ]
    ).count()
    waiting_documents = required_documents.filter(status='waiting').count()

    context = {
        'client_case': client_case,
        'required_documents': required_documents,
        'total_documents': total_documents,
        'completed_documents': completed_documents,
        'waiting_documents': waiting_documents,
    }

    return render(request, 'documents/staff_client_case_detail.html', context)

@login_required
def feedback_create(request):
    if request.method == 'POST':
        form = FeedbackForm(request.POST)

        if form.is_valid():
            feedback = form.save(commit=False)
            feedback.user = request.user
            feedback.save()

            if request.user.is_staff:
                return redirect('staff_dashboard')

            return redirect('document_list')
    else:
        form = FeedbackForm()

    return render(request, 'documents/feedback_form.html', {
        'form': form
    })